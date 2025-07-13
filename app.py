import os
import time
import math
import re
import logging
from datetime import datetime, timedelta
from functools import wraps
from contextlib import contextmanager
from flask import (
    Flask, request, render_template, redirect,
    url_for, session, send_file, flash
)
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from Crypto.Util.Padding import pad
from werkzeug.utils import secure_filename
from PIL import Image
import sqlite3
from cryptography.fernet import Fernet
from apscheduler.schedulers.background import BackgroundScheduler
import bcrypt
import shutil
import gzip
import smtplib
from email.mime.text import MIMEText
import bleach
from forms import (
    RegistrationForm, LoginForm,
    ContactForm, UploadForm, OptionsForm
)
from werkzeug.security import generate_password_hash, check_password_hash
from flask_wtf.csrf import CSRFProtect
from flask_mail import Mail, Message
from matplotlib import pyplot as plt
import io
import base64
from matplotlib.gridspec import GridSpec
import matplotlib
matplotlib.use('Agg')
import numpy as np
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    Table, TableStyle
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import pyotp


# Logging Setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# Flask App Setup
app = Flask(__name__)
app.secret_key = os.urandom(24)  # For production, use a fixed secret in an env var

app.config.update({
    'WTF_CSRF_ENABLED': True,
    'MAX_CONTENT_LENGTH': 50 * 1024 * 1024,  # 50 MB
    'UPLOAD_FOLDER': 'static/uploads',
    'ALLOWED_EXTENSIONS': {'png', 'jpg', 'jpeg', 'bmp', 'gif'},
    'TEMPLATES_AUTO_RELOAD': True,
    
    # Mail config — replaced sensitive data
    'MAIL_SERVER': 'smtp.gmail.com',
    'MAIL_PORT': 587,
    'MAIL_USE_TLS': True,
    'MAIL_USERNAME': 'your_email@example.com',  
    'MAIL_PASSWORD': 'your_email_password',    
    'MAIL_DEFAULT_SENDER': ('Your App', 'your_email@example.com'),

    # Security/session
    'SESSION_COOKIE_SECURE': True,
    'SESSION_COOKIE_HTTPONLY': True,
    'SESSION_COOKIE_SAMESITE': 'Lax',
    'PERMANENT_SESSION_LIFETIME': 1800,
    'WTF_CSRF_TIME_LIMIT': 1800,
    
    # App-specific config
    'DATABASE': 'data/app.db',
    'BACKUP_DIR': 'backups',
    'MAX_BACKUPS': 30,
    'BACKUP_ENCRYPTION_KEY': Fernet.generate_key(),
    'ADMIN_EMAIL': 'admin@example.com'  
})



# Initialize Extensions
csrf = CSRFProtect(app)
mail = Mail(app)

limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)



# Directory Setup
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])
    logger.info(f"Created upload directory at {app.config['UPLOAD_FOLDER']}")
elif not os.access(app.config['UPLOAD_FOLDER'], os.W_OK):
    logger.error(f"Upload directory not writable: {app.config['UPLOAD_FOLDER']}")

os.makedirs('data', exist_ok=True)
os.makedirs(app.config['BACKUP_DIR'], exist_ok=True)



# Security Event Logger
def log_security_event(event_type, description, user=None, ip=None):
    ip = ip or request.remote_addr
    user = user or session.get('username')

    logger.warning(
        f"SECURITY EVENT - {event_type}: {description}. "
        f"User: {user}, IP: {ip}, Path: {request.path}"
    )

    if event_type in ['failed_login', 'brute_force', 'xss_attempt']:
        try:
            msg = Message(
                subject=f"Security Alert: {event_type}",
                recipients=[app.config['MAIL_USERNAME']],
                body=f"""
                Security Alert:
                Type: {event_type}
                Description: {description}
                User: {user}
                IP: {ip}
                Path: {request.path}
                Time: {datetime.now()}
                """
            )
            mail.send(msg)
        except Exception as e:
            logger.error(f"Failed to send security alert email: {e}")



# Database Manager
class DatabaseManager:
    def __init__(self, app=None):
        self.app = app
        if app is not None:
            self.init_app(app)

    def init_app(self, app):
        self.app = app
        self._verify_encryption_key()
        self.init_databases()

    def _verify_encryption_key(self):
        try:
            Fernet(self.app.config['BACKUP_ENCRYPTION_KEY'])
        except ValueError:
            logger.warning("Invalid encryption key, generating new one")
            self.app.config['BACKUP_ENCRYPTION_KEY'] = Fernet.generate_key()

    @contextmanager
    def get_db(self):
        db_path = self.app.config['DATABASE']
        conn = sqlite3.connect(db_path, timeout=20)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
            conn.commit()
        except Exception as e:
            conn.rollback()
            logger.error(f"Database error: {str(e)}")
            raise
        finally:
            conn.close()

    def init_databases(self):
        try:
            with self.get_db() as conn:
                conn.executescript('''
                    CREATE TABLE IF NOT EXISTS users (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        username TEXT UNIQUE NOT NULL,
                        hashed_password TEXT NOT NULL,
                        login_attempts INTEGER DEFAULT 0,
                        last_attempt TIMESTAMP,
                        totp_secret TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    );
                    CREATE INDEX IF NOT EXISTS idx_users_username ON users (username);

                    CREATE TABLE IF NOT EXISTS uploads (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id INTEGER,
                        filename TEXT NOT NULL,
                        mode TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        FOREIGN KEY(user_id) REFERENCES users(id)
                    );
                    CREATE INDEX IF NOT EXISTS idx_uploads_user_id ON uploads (user_id);

                    CREATE TABLE IF NOT EXISTS system_metrics (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        total_users INTEGER,
                        active_users INTEGER,
                        total_uploads INTEGER,
                        total_storage INTEGER
                    );
                ''')
            logger.info("Database tables created/verified successfully.")
        except sqlite3.Error as e:
            logger.error(f"Database initialization error: {e}")
    def create_backup(self):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(self.app.config['BACKUP_DIR'], f"backup_{timestamp}")

        try:
            os.makedirs(backup_path)

            # Backup the database
            db_path = self.app.config['DATABASE']
            with open(db_path, 'rb') as f_in:
                with gzip.open(os.path.join(backup_path, "app.db.gz"), 'wb') as f_out:
                    shutil.copyfileobj(f_in, f_out)

            # Verify the backup
            if not self.verify_backup(os.path.join(backup_path, "app.db.gz")):
                logger.error(f"Backup verification failed for: {backup_path}")
                raise Exception("Backup verification failed")

            self.rotate_backups()
            return backup_path
        except Exception as e:
            logger.error(f"Backup failed: {str(e)}")
            raise

    def verify_backup(self, backup_file):
        """Verifies the integrity of the backup by attempting to read its contents."""
        try:
            with gzip.open(backup_file, 'rb') as f_in:
                f_in.read(1024)  # Read small chunk to verify
            logger.info(f"Backup verified successfully: {backup_file}")
            return True
        except Exception as e:
            logger.error(f"Backup verification failed: {e}")
            return False

    def rotate_backups(self):
        backups = sorted(
            [d for d in os.scandir(self.app.config['BACKUP_DIR']) if d.is_dir()],
            key=lambda x: x.stat().st_mtime,
            reverse=True
        )
        for backup in backups[self.app.config['MAX_BACKUPS']:]:
            shutil.rmtree(backup.path)

    def get_system_stats(self):
        stats = {}
        try:
            with self.get_db() as conn:
                stats['total_users'] = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
                stats['active_users'] = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]  # Placeholder logic
                stats['total_uploads'] = conn.execute("SELECT COUNT(*) FROM uploads").fetchone()[0]

                # Calculate total storage
                stats['total_storage'] = 0
                for row in conn.execute("SELECT filename FROM uploads"):
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], row['filename'])
                    try:
                        stats['total_storage'] += os.path.getsize(filepath)
                    except FileNotFoundError:
                        logger.warning(f"File not found: {filepath}")

        except sqlite3.Error as e:
            logger.error(f"Error retrieving system stats: {e}")
            return {}
        return stats


def init_db():
    conn = sqlite3.connect('messages.db')  
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS messages (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT,
                    email TEXT,
                    subject TEXT,
                    message TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )''')
    conn.commit()
    conn.close()



# Backup Scheduler
class BackupScheduler:
    def __init__(self, db_manager):
        self.db_manager = db_manager
        self.scheduler = BackgroundScheduler()
        self.logger = logging.getLogger(__name__)

    def start(self):
        try:
            # Daily backup at 2 AM
            self.scheduler.add_job(
                self.run_backup,
                'cron',
                hour=2,
                minute=0
            )
            # Weekly metrics on Sundays at 3 AM
            self.scheduler.add_job(
                self.record_metrics,
                'cron',
                day_of_week='sun',
                hour=3
            )
            self.scheduler.start()
        except Exception as e:
            self.logger.error(f"Scheduler error: {str(e)}")
            raise

    def run_backup(self):
        try:
            path = self.db_manager.create_backup()
            self.logger.info(f"Backup created: {path}")
        except Exception as e:
            self.logger.error(f"Backup failed: {str(e)}")

    def record_metrics(self):
        try:
            stats = self.db_manager.get_system_stats()
            with self.db_manager.get_db() as conn:
                conn.execute('''
                    INSERT INTO system_metrics 
                    (total_users, active_users, total_uploads, total_storage) 
                    VALUES (?, ?, ?, ?)
                ''', (
                    stats.get('total_users', 0),
                    stats.get('active_users', 0),
                    stats.get('total_uploads', 0),
                    stats.get('total_storage', 0)
                ))
            logger.info("System metrics recorded successfully.")
        except sqlite3.Error as e:
            logger.error(f"Error recording system metrics: {e}")



# Initialize services
db_manager = DatabaseManager(app)

if not app.debug: n'
    backup_scheduler = BackupScheduler(db_manager)
    backup_scheduler.start()



# Helper Functions
def get_upload_path(filename):
    """Get absolute path for uploaded files with proper sanitization"""
    safe_filename = sanitize_filename_path(filename) 
    return os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], safe_filename))
def generate_text_report(report_data):
    """Generate formatted text report"""
    lines = [
        f"{'=' * 50}",
        f"{report_data['header']['title']:^50}",
        f"{'=' * 50}",
        f"Date: {report_data['header']['date']}",
        f"Author: {report_data['header']['author']}",
        f"File: {report_data['header']['filename']}",
        "",
        "SUMMARY",
        "-" * 50,
        report_data['summary']['overview'],
        "",
        "RESULTS",
        "-" * 50
    ]

    for result in report_data['results']:
        lines.extend([
            f"\nMode: {result['mode']}",
            f"Security: {result['security']}",
            f"IV: {result['iv']}",
            f"Integrity: {result['integrity']}",
            f"Time: {result['time']}",
            f"Description: {result['description']}"
        ])

    lines.extend([
        "",
        "CONCLUSION",
        "-" * 50,
        report_data['summary']['conclusion'],
        "",
        f"{'=' * 50}",
        "End of Report"
    ])

    return "\n".join(lines)


def generate_pdf_report(report_data, buffer):
    try:
        # Register custom fonts (if available)
        font_dir = os.path.join(app.static_folder, 'fonts')
        professional_fonts = [
            ('Montserrat-Regular', 'Montserrat-Regular.ttf'),
            ('Roboto-Regular', 'Roboto-Regular.ttf'),
            ('OpenSans-Regular', 'OpenSans-Regular.ttf'),
            ('Lato-Regular', 'Lato-Regular.ttf')
        ]

        for font_name, font_file in professional_fonts:
            try:
                pdfmetrics.registerFont(TTFont(font_name, os.path.join(font_dir, font_file)))
            except:
                logger.warning(f"Could not register font {font_file}")

        # Define color scheme
        primary_color = HexColor('#2c3e50')
        secondary_color = HexColor('#3498db')
        accent_color = HexColor('#e74c3c')

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            name='ReportTitle',
            fontName='Montserrat-Regular',
            fontSize=18,
            leading=22,
            alignment=1,
            spaceAfter=20,
            textColor=primary_color
        )
        header_style = ParagraphStyle(
            name='SectionHeader',
            fontName='Montserrat-Regular',
            fontSize=14,
            leading=18,
            spaceBefore=20,
            spaceAfter=10,
            textColor=secondary_color
        )
        normal_style = ParagraphStyle(
            name='NormalText',
            fontName='Roboto-Regular',
            fontSize=10,
            leading=12,
            spaceAfter=12,
            textColor=HexColor('#333333')
        )

        # Build PDF
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []

        story.append(Paragraph(report_data['header']['title'], title_style))
        story.append(Spacer(1, 12))

        meta_text = f"""
        <b>Date:</b> {report_data['header']['date']}<br/>
        <b>Author:</b> {report_data['header']['author']}<br/>
        <b>File:</b> {report_data['header']['filename']}
        """
        story.append(Paragraph(meta_text, normal_style))
        story.append(Spacer(1, 24))

        story.append(Paragraph('SUMMARY', header_style))
        story.append(Paragraph(report_data['summary']['overview'], normal_style))
        story.append(Spacer(1, 24))

        story.append(Paragraph('ENCRYPTION MODE COMPARISON', header_style))

        for result in report_data['results']:
            mode_text = f"""
            <b>{result['mode']} MODE</b><br/>
            <b>Security:</b> {result['security']}<br/>
            <b>IV:</b> {result['iv']}<br/>
            <b>Integrity:</b> {result['integrity']}<br/>
            <b>Time:</b> {result['time']}<br/>
            <b>Description:</b> {result['description']}
            """
            story.append(Paragraph(mode_text, normal_style))
            story.append(Spacer(1, 16))

        story.append(Paragraph('CONCLUSION', header_style))
        story.append(Paragraph(report_data['summary']['conclusion'], normal_style))

        doc.build(story)
        return True

    except Exception as e:
        logger.error(f"Error in generate_pdf_report: {str(e)}", exc_info=True)
        return False


def generate_word_report(report_data, filename):
    """Generate professional Word report"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading(report_data['header']['title'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Date: {report_data['header']['date']}")
        doc.add_paragraph(f"Author: {report_data['header']['author']}")
        doc.add_paragraph(f"File: {report_data['header']['filename']}")
        doc.add_paragraph()

        doc.add_heading('Summary', level=2)
        doc.add_paragraph(report_data['summary']['overview'])
        doc.add_paragraph()

        doc.add_heading('Encryption Mode Comparison', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Mode'
        hdr_cells[1].text = 'Security'
        hdr_cells[2].text = 'IV'
        hdr_cells[3].text = 'Integrity'
        hdr_cells[4].text = 'Time'
        hdr_cells[5].text = 'Description'

        for result in report_data['results']:
            row_cells = table.add_row().cells
            row_cells[0].text = result['mode']
            row_cells[1].text = result['security']
            row_cells[2].text = result['iv']
            row_cells[3].text = result['integrity']
            row_cells[4].text = result['time']
            row_cells[5].text = result['description']

        doc.add_paragraph()

        doc.add_heading('Visual Comparison', level=2)

        # Original Image
        original_path = os.path.join(app.config['UPLOAD_FOLDER'], report_data['visuals']['original_img'])
        if os.path.exists(original_path):
            doc.add_heading('Original Image', level=3)
            doc.add_picture(original_path, width=Inches(2.5))

        # Encrypted Images
        for mode in ['ECB', 'CBC', 'GCM']:
            img_key = f"{mode.lower()}_img"
            img_path = os.path.join(app.config['UPLOAD_FOLDER'], report_data['visuals'].get(img_key, ''))
            if os.path.exists(img_path):
                doc.add_heading(f'{mode} Mode', level=3)
                doc.add_picture(img_path, width=Inches(2.5))

        # Comparison Chart
        chart_path = os.path.join(app.config['UPLOAD_FOLDER'], report_data['visuals']['comparison_img'])
        if os.path.exists(chart_path):
            doc.add_heading('Performance Comparison', level=2)
            doc.add_picture(chart_path, width=Inches(6))

        doc.add_heading('Conclusion', level=2)
        doc.add_paragraph(report_data['summary']['conclusion'])

        doc.save(filename)
        return True

    except Exception as e:
        logger.error(f"Error generating Word report: {e}")
        return False

def save_image(img_array, filename):
    """Saves a NumPy array as a PNG image."""
    # Convert numpy array to PIL Image object
    img = Image.fromarray(img_array)  
    img.save(filename)


def ciphertext_to_figure(ct_bytes):
    """Convert ciphertext bytes to a 2D numpy array suitable for matplotlib display."""
    width = 256
    length = len(ct_bytes)
    height = math.ceil(length / width)
    padded_len = width * height
    # Pad ciphertext bytes with null bytes so total length fits height*width
    padded_bytes = ct_bytes + b'\x00' * (padded_len - length)
    # Convert bytes to uint8 numpy array and reshape to 2D
    return np.frombuffer(padded_bytes, dtype=np.uint8).reshape(height, width)


def allowed_file(filename):
    """Check if the filename has an allowed extension and matches a safe pattern."""
    # Ensure filename matches pattern: letters, numbers, underscore, space, dash, with extension 3-4 chars
    if not re.match(r'^[\w\s-]+\.[A-Za-z]{3,4}$', filename):
        return False
    # Check extension is in allowed list
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def sanitize_input(input_string):
    """Sanitize user input to prevent XSS and SQL injection attacks."""
    if not input_string:
        return input_string

    # Remove potentially dangerous characters commonly used in injections
    input_string = re.sub(r'[;\\\'\"\0\x08\x09\x0a\x0b\x0c\x0d\x1a]', '', input_string)

    # Use bleach to strip any remaining HTML tags safely
    cleaned = bleach.clean(input_string, strip=True)

    # Additional regex patterns to detect SQL injection attempts
    sql_injection_patterns = [
        r'\b(ALTER|CREATE|DELETE|DROP|EXEC(UTE){0,1}|INSERT( +INTO){0,1}|MERGE|SELECT|UPDATE|UNION( +ALL){0,1})\b',
        r'\b(OR\s+\d+\s*=\s*\d+)\b',
        r'(\-\-|\#|\/\*[\s\S]*?\*\/)'
    ]

    for pattern in sql_injection_patterns:
        if re.search(pattern, cleaned, re.IGNORECASE):
            raise ValueError("Invalid input detected")

    return cleaned


def sanitize_filename_path(filename):
    """Sanitize filename to prevent directory traversal and unsafe characters."""
    filename = filename.replace('\\', '/')
    # Remove any patterns trying to navigate up directories
    filename = re.sub(r'(\.\./|\.\.\\|\.\.\/|\.\\).*', '', filename)
    # Remove relative path indicators
    filename = re.sub(r'(\./|\.\\|\.\/|\.\\).*', '', filename)
    # Get base filename only
    filename = os.path.basename(filename)
    # Secure the filename with Werkzeug's secure_filename, fallback to default name if empty
    filename = secure_filename(filename) or "default_filename.txt"
    return filename


def encrypt_ecb(data, key):
    """Encrypt data using AES ECB mode."""
    cipher = AES.new(key, AES.MODE_ECB)
    padded = pad(data, AES.block_size)
    start = time.time()
    ct = cipher.encrypt(padded)
    # Return ciphertext, IV (None for ECB), tag (None for ECB), and elapsed time
    return ct, None, None, time.time() - start


def encrypt_cbc(data, key):
    """Encrypt data using AES CBC mode."""
    iv = get_random_bytes(AES.block_size)  # Generate random IV
    cipher = AES.new(key, AES.MODE_CBC, iv)
    padded = pad(data, AES.block_size)
    start = time.time()
    ct = cipher.encrypt(padded)
    # Return ciphertext, IV, no tag for CBC, and elapsed time
    return ct, iv, None, time.time() - start


def encrypt_gcm(data, key):
    """Encrypt data using AES GCM mode with authentication tag."""
    iv = get_random_bytes(12)  # GCM standard nonce size is 12 bytes
    cipher = AES.new(key, AES.MODE_GCM, nonce=iv)
    start = time.time()
    ct, tag = cipher.encrypt_and_digest(data)
    # Return ciphertext, IV, authentication tag, and elapsed time
    return ct, iv, tag, time.time() - start


def security_level(mode):
    """Return a user-friendly string describing the security level of the AES mode."""
    levels = {
        "ECB": "🟡 Low - Pattern leakage, not semantically secure",
        "CBC": "🟠 Medium - Requires IV, better than ECB but no integrity",
        "GCM": "🟢 High - Confidentiality + integrity checks"
    }
    return levels.get(mode, "⚪ Unknown")


def integrity_check(mode):
    """Return if the AES mode provides integrity checking."""
    return "✅ Yes (via authentication tag)" if mode == "GCM" else "❌ No"


def recommend_best_mode():
    """Recommendation string favoring GCM for best security and performance."""
    return "🌟 Recommendation: GCM provides the best security with integrity checks and good performance"


def ciphertext_to_png(ct_bytes, filename):
    """Convert ciphertext bytes to grayscale PNG image."""
    width = 256
    length = len(ct_bytes)
    height = math.ceil(length / width)
    padded_len = width * height
    # Pad bytes to fill the image grid exactly
    padded_bytes = ct_bytes + b'\x00' * (padded_len - length)
    # Create grayscale image from padded bytes
    img = Image.frombytes('L', (width, height), padded_bytes)
    img.save(filename)


# Authentication Decorator
def login_required(f):
    """Decorator to ensure user is logged in before accessing a route."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            flash('🔒 You must be logged in to access this page', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)

    return decorated_function


def get_user_id(username):
    """Lookup user ID from the database given a username."""
    with db_manager.get_db() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT id FROM users WHERE username = ?', (username,))
        user = cursor.fetchone()
        return user['id'] if user else None


# Error handlers

@app.errorhandler(400)
def bad_request(e):
    """Handle 400 Bad Request errors by logging and rendering a friendly page."""
    logger.warning(f"400 Bad Request: {request.url} - {str(e)}")
    return render_template('400.html'), 400


@app.errorhandler(404)
def page_not_found(e):
    """Handle 404 Not Found errors."""
    return render_template('404.html'), 404


@app.errorhandler(413)
def file_too_large(e):
    """Handle 413 Payload Too Large errors with user-friendly flash message."""
    flash('⚠️ File too large! Maximum size is 50MB', 'error')
    return redirect(request.url)


@app.errorhandler(500)
def internal_error(e):
    """Handle 500 Internal Server Error."""
    return render_template('500.html'), 500


# Initialize the database table when app context is active
with app.app_context():
    init_db()


@app.route('/')
def home():
    """Landing page. Redirects logged-in users to upload page, else shows index."""
    if 'username' in session:
        return redirect(url_for('upload'))
    return render_template('index.html')

@app.route('/email', methods=['GET', 'POST'])
def email():
    """Handles the contact form submission: saves messages to DB and sends email notification."""
    img = "/static/images/logo.png"   # Logo image path for template
    noty = "/static/sounds/notification.mp3"  # Notification sound path for template
    form = ContactForm()  # Instantiate contact form

    if form.validate_on_submit():
        # Sanitize all user inputs to prevent XSS/SQL injection
        name = sanitize_input(form.name.data)
        email = sanitize_input(form.email.data)
        subject = sanitize_input(form.subject.data)
        message = sanitize_input(form.message.data)

        # Save message to local SQLite database
        try:
            conn = sqlite3.connect('messages.db')
            c = conn.cursor()
            c.execute("INSERT INTO messages (name, email, subject, message) VALUES (?, ?, ?, ?)",
                      (name, email, subject, message))
            conn.commit()
            flash('Your message has been saved successfully!', 'success')
        except sqlite3.Error as e:
            # Handle database errors gracefully and log them
            flash(f'Error saving message to database: {e}', 'error')
            logger.error(f"Database error: {str(e)}")
        finally:
            if conn:
                conn.close()

        # Attempt to send an email notification about the new message
        try:
            send_email_notification(name, email, subject, message)
            flash('Email notification sent successfully!', 'success')
        except Exception as e:
            # If email sending fails, log error and notify user
            logger.error(f"Email failed: {str(e)}")
            flash('Email notification failed to send. Message was saved locally.', 'warning')
    else:
        # If form validation failed, flash errors for each field
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"{getattr(form, field).label.text}: {error}", 'error')

    # Render the email/contact page template with form and media assets
    return render_template('email.html', img=img, noty=noty, form=form)


def send_email_notification(name, email, subject, message):
    """Sends an email notification about a new contact form submission."""
    try:
        # Validate necessary email config is present
        if not all([app.config['MAIL_USERNAME'], app.config['MAIL_PASSWORD']]):
            raise ValueError("Email configuration incomplete")

        # Compose the email message
        msg = Message(
            subject=f"New Contact: {subject}",
            recipients=[app.config['MAIL_USERNAME']],  # Send notification to configured email
            body=f"""
            New Contact Form Submission:
            From: {name} <{email}>
            Subject: {subject}
            Message: 
            {message}
            """
        )

        logger.info(f"Attempting to send email via {app.config['MAIL_SERVER']}")

        # Use Flask-Mail connection context manager to send the message
        with mail.connect() as conn:
            if not conn:
                raise ConnectionError("Could not connect to SMTP server")

            conn.send(msg)

        logger.info("Email sent successfully")
        return True

    except Exception as e:
        # Log detailed exception info for debugging
        logger.error(f"Email failed: {str(e)}", exc_info=True)
        # Raise an error for the caller to handle
        raise RuntimeError(f"Could not send email: {str(e)}") from e


@app.route('/learn')
def learn():
    """Render a static 'learn' page."""
    return render_template('learn.html')


@app.route('/register', methods=['GET', 'POST'])
@limiter.limit("5 per hour")  # Rate limit registration to 5 per hour per IP/user
def register():
    """Handle user registration with password complexity and 2FA secret generation."""
    if 'username' in session:
        return redirect(url_for('upload'))  # Redirect logged-in users to upload page

    form = RegistrationForm()
    if form.validate_on_submit():
        username = sanitize_input(form.username.data)
        password = form.password.data

        # Enforce strong password policy: length >=12, upper, lower, digit, special char
        if not (len(password) >= 12 and
                any(c.isupper() for c in password) and
                any(c.islower() for c in password) and
                any(c.isdigit() for c in password) and
                any(not c.isalnum() for c in password)):
            flash('Password must be at least 12 characters with uppercase, lowercase, number, and special character',
                  'error')
            return redirect(url_for('register'))

        # Hash password with bcrypt (cost factor 14)
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt(rounds=14))

        try:
            with db_manager.get_db() as conn:
                # Generate a secret key for 2FA (TOTP)
                totp_secret = pyotp.random_base32()

                # Insert new user into the database
                conn.execute(
                    'INSERT INTO users (username, hashed_password, totp_secret) VALUES (?, ?, ?)',
                    (username, hashed_password.decode('utf-8'), totp_secret))

                flash('🎉 Registration successful! Please log in', 'success')
                return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            # Handle case where username already exists
            flash('⚠️ Username already exists', 'error')
        except sqlite3.Error as e:
            logger.error(f"Database error during registration: {e}")
            flash('⚠️ Database error occurred. Please try again.', 'error')

    else:
        # Flash any form validation errors
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"{getattr(form, field).label.text}: {error}", 'error')

    # Render registration page with form
    return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")  # Rate limit login attempts
def login():
    """Authenticate user, enforce lockout policy, and establish session on success."""
    if 'username' in session:
        return redirect(url_for('upload'))

    form = LoginForm()
    if form.validate_on_submit():
        username = sanitize_input(form.username.data)
        password = form.password.data

        try:
            with db_manager.get_db() as conn:
                cursor = conn.cursor()
                cursor.execute('''SELECT id, username, hashed_password, login_attempts, last_attempt, totp_secret
                                 FROM users WHERE username = ?''', (username,))
                user = cursor.fetchone()

                if user:
                    # Lock account if too many failed attempts within 30 minutes
                    if user['login_attempts'] >= 5 and user['last_attempt'] and (
                            datetime.now() - datetime.strptime(user['last_attempt'], '%Y-%m-%d %H:%M:%S') < timedelta(
                        minutes=30)):
                        flash('🔒 Account temporarily locked due to too many failed attempts. Try again later.', 'error')
                        return redirect(url_for('login'))

                    # Verify password using bcrypt
                    if bcrypt.checkpw(password.encode('utf-8'), user['hashed_password'].encode('utf-8')):
                        # Reset login attempts on successful login
                        cursor.execute('''UPDATE users SET login_attempts = 0, last_attempt = NULL
                                          WHERE username = ?''', (username,))
                        conn.commit()

                        # Set user session
                        session['username'] = username
                        session.permanent = True
                        flash(f'👋 Welcome back, {username}!', 'success')
                        return redirect(url_for('upload'))
                    else:
                        # Increment failed login attempts and record timestamp
                        cursor.execute('''UPDATE users
                                          SET login_attempts = login_attempts + 1,
                                              last_attempt = datetime('now')
                                          WHERE username = ?''', (username,))
                        conn.commit()
                        log_security_event('failed_login', f'Failed login attempt for user {username}', username=username)
                        flash('🔐 Invalid username or password', 'error')
                else:
                    # Log attempts for usernames that don't exist
                    log_security_event('invalid_user', f'Login attempt with invalid username {username}', username=username)
                    flash('🔐 Invalid username or password', 'error')
        except sqlite3.Error as e:
            logger.error(f"Database error during login: {e}")
            flash('⚠️ Database error occurred. Please try again.', 'error')

    else:
        # Flash any form validation errors
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"{getattr(form, field).label.text}: {error}", 'error')

    # Render login page with form
    return render_template('login.html', form=form)


@app.route('/logout')
def logout():
    """Log out the current user by clearing session and redirecting to login."""
    session.pop('username', None)
    flash('👋 Logged out successfully!', 'info')
    return redirect(url_for('login'))

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    """
    Handles secure file upload for authenticated users.
    - Validates file type and sanitizes filename.
    - Stores file securely with a timestamp prefix.
    - Records upload details in the database.
    - Stores path in session for later processing.
    """
    form = UploadForm()

    if form.validate_on_submit():
        file = form.file.data
        if file and allowed_file(file.filename):
            # Sanitize filename to prevent path traversal
            filename = sanitize_filename_path(file.filename)

            # Append timestamp to filename for uniqueness
            unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            file.save(filepath)  # Save the file

            # Fetch user ID based on session username
            user_id = get_user_id(session['username'])
            if user_id:
                try:
                    with db_manager.get_db() as conn:
                        # Record file upload in DB
                        conn.execute(
                            'INSERT INTO uploads (user_id, filename, mode) VALUES (?, ?, ?)',
                            (user_id, unique_filename, None)
                        )
                    flash('📁 File uploaded successfully!', 'success')
                    session['uploaded_file'] = filepath
                    return redirect(url_for('options'))
                except sqlite3.Error as e:
                    logger.error(f"Database error during file upload: {e}")
                    flash('⚠️ Database error occurred. Please try again.', 'error')
            else:
                flash('⚠️ User ID not found. Please log in again.', 'error')
                return redirect(url_for('login'))
        else:
            flash('⚠️ Invalid file type or filename. Allowed: PNG, JPG, JPEG, BMP, GIF', 'error')
    else:
        # Flash WTForms validation errors
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"{getattr(form, field).label.text}: {error}", 'error')

    return render_template('upload.html', form=form)


@app.route('/options', methods=['GET', 'POST'])
@login_required
def options():
    """
    Displays available options after image upload:
    - Compare encryption modes.
    - Encrypt using a selected mode (ECB, CBC, GCM).
    """
    if 'uploaded_file' not in session:
        flash('⚠️ Please upload an image first', 'warning')
        return redirect(url_for('upload'))

    form = OptionsForm()

    if request.method == 'POST':
        action = request.form.get('action')
        mode = request.form.get('mode')

        if action == 'compare':
            return redirect(url_for('compare'))

        elif action == 'encrypt':
            if mode not in ['ECB', 'CBC', 'GCM']:
                flash('⚠️ Invalid encryption mode', 'error')
                return redirect(url_for('options'))
            return redirect(url_for('encrypt_image', mode=mode))

        else:
            flash('⚠️ Invalid action', 'error')
            return redirect(url_for('options'))

    return render_template('options.html', form=form)


@app.route('/compare')
@login_required
def compare():
    """
    Compares ECB, CBC, and GCM AES encryption modes:
    - Encrypts uploaded image using each mode.
    - Visualizes encryption effects and timing.
    - Displays side-by-side chart and analysis.
    - Prepares report data for download.
    """
    if 'uploaded_file' not in session:
        flash('⚠️ Please upload an image first', 'warning')
        return redirect(url_for('upload'))

    filepath = session['uploaded_file']
    if not os.path.exists(filepath):
        flash(f'⚠️ File not found at path: {filepath}. Please upload again.', 'error')
        session.pop('uploaded_file', None)
        return redirect(url_for('upload'))

    try:
        with open(filepath, 'rb') as f:
            data = f.read()
    except Exception as e:
        flash(f'⚠️ Error reading file: {str(e)}', 'error')
        session.pop('uploaded_file', None)
        return redirect(url_for('upload'))

    key = get_random_bytes(16)  # Generate AES-128 key

    # --- Encrypt using all 3 modes and record time ---
    ct_ecb, iv_ecb, tag_ecb, time_ecb = encrypt_ecb(data, key)
    ct_cbc, iv_cbc, tag_cbc, time_cbc = encrypt_cbc(data, key)
    ct_gcm, iv_gcm, tag_gcm, time_gcm = encrypt_gcm(data, key)

    # --- Visualization using matplotlib ---
    fig = plt.figure(figsize=(15, 10), dpi=100)
    gs = GridSpec(2, 4, figure=fig)

    # Original image
    ax1 = fig.add_subplot(gs[0, 0])
    try:
        original_pil_img = Image.open(filepath)
        ax1.imshow(original_pil_img)
    except Exception as e:
        logger.error(f"Could not open original image for plot: {filepath} - {e}")
        ax1.text(0.5, 0.5, 'Error loading original image', ha='center', va='center')
    ax1.set_title('Original Image', fontsize=10)
    ax1.axis('off')

    # ECB visualization
    ax2 = fig.add_subplot(gs[0, 1])
    ax2.imshow(ciphertext_to_figure(ct_ecb), cmap='gray')
    ax2.set_title('ECB Mode\n(PATTERN LEAKAGE)', fontsize=10, color='red')
    ax2.axis('off')

    # CBC visualization
    ax3 = fig.add_subplot(gs[0, 2])
    ax3.imshow(ciphertext_to_figure(ct_cbc), cmap='gray')
    ax3.set_title('CBC Mode\n(Secure)', fontsize=10, color='green')
    ax3.axis('off')

    # GCM visualization
    ax4 = fig.add_subplot(gs[0, 3])
    ax4.imshow(ciphertext_to_figure(ct_gcm), cmap='gray')
    ax4.set_title('GCM Mode\n(Best Security)', fontsize=10, color='blue')
    ax4.axis('off')

    # --- Performance Chart ---
    ax5 = fig.add_subplot(gs[1, :])
    modes = ['ECB', 'CBC', 'GCM']
    times_ms = [time_ecb * 1000, time_cbc * 1000, time_gcm * 1000]
    colors = ['red', 'green', 'blue']
    ax5.bar(modes, times_ms, color=colors)
    ax5.set_ylabel('Encryption Time (ms)')
    ax5.set_title('Performance Comparison')
    ax5.set_ylim(0, max(times_ms) * 1.1 if max(times_ms) > 0 else 1)

    # Save comparison figure
    comparison_filename_base = f"comparison_{datetime.now().strftime('%Y%m%d%H%M%S')}.png"
    comparison_figure_path = os.path.join(app.config['UPLOAD_FOLDER'], comparison_filename_base)
    try:
        plt.tight_layout()
        plt.savefig(comparison_figure_path, bbox_inches='tight')
    except Exception as e:
        logger.error(f"Error saving comparison figure: {e}")
        flash("Error generating comparison figure.", "error")
    finally:
        plt.close(fig)

    # --- Save encrypted versions ---
    save_image(ciphertext_to_figure(ct_ecb), os.path.join(app.config['UPLOAD_FOLDER'],
               f"ecb_encrypted_{datetime.now().strftime('%Y%m%d%H%M%S')}.png"))
    save_image(ciphertext_to_figure(ct_cbc), os.path.join(app.config['UPLOAD_FOLDER'],
               f"cbc_encrypted_{datetime.now().strftime('%Y%m%d%H%M%S')}.png"))
    save_image(ciphertext_to_figure(ct_gcm), os.path.join(app.config['UPLOAD_FOLDER'],
               f"gcm_encrypted_{datetime.now().strftime('%Y%m%d%H%M%S')}.png"))

    # --- Prepare results for template ---
    results = {
        'ECB': {
            "Security": security_level("ECB"),
            "IV": iv_ecb.hex()[:16] + "..." if iv_ecb else "None",
            "Integrity": integrity_check("ECB"),
            "Time (ms)": f"{time_ecb * 1000:.2f}",
            "color": "warning"
        },
        'CBC': {
            "Security": security_level("CBC"),
            "IV": iv_cbc.hex()[:16] + "..." if iv_cbc else "None",
            "Integrity": integrity_check("CBC"),
            "Time (ms)": f"{time_cbc * 1000:.2f}",
            "color": "info"
        },
        'GCM': {
            "Security": security_level("GCM"),
            "IV": iv_gcm.hex()[:16] + "..." if iv_gcm else "None",
            "Integrity": integrity_check("GCM"),
            "Time (ms)": f"{time_gcm * 1000:.2f}",
            "color": "success"
        }
    }

    # --- Build report metadata ---
    report_content = {
        "header": {
            "title": "AES Encryption Mode Analysis Report",
            "date": datetime.now().strftime('%B %d, %Y'),
            "author": session.get('username', 'Anonymous User'),
            "filename": os.path.basename(filepath)
        },
        "summary": {
            "overview": "This report compares the performance and security characteristics of three AES encryption modes: ECB, CBC, and GCM.",
            "conclusion": recommend_best_mode()
        },
        "results": [
            {
                "mode": "ECB",
                "security": security_level("ECB"),
                "iv": iv_ecb.hex()[:16] + "...",
                "integrity": integrity_check("ECB"),
                "time": f"{time_ecb * 1000:.2f} ms",
                "description": "ECB is the simplest encryption mode, but leaks patterns due to lack of IV."
            },
            {
                "mode": "CBC",
                "security": security_level("CBC"),
                "iv": iv_cbc.hex()[:16] + "...",
                "integrity": integrity_check("CBC"),
                "time": f"{time_cbc * 1000:.2f} ms",
                "description": "CBC uses an IV to randomize ciphertext, improving security over ECB."
            },
            {
                "mode": "GCM",
                "security": security_level("GCM"),
                "iv": iv_gcm.hex()[:16] + "...",
                "integrity": integrity_check("GCM"),
                "time": f"{time_gcm * 1000:.2f} ms",
                "description": "GCM provides authenticated encryption, making it the most secure option."
            }
        ],
        "visuals": {
            "original_img": os.path.basename(filepath),
            "ecb_img": f"ecb_encrypted_{datetime.now().strftime('%Y%m%d%H%M%S')}.png",
            "cbc_img": f"cbc_encrypted_{datetime.now().strftime('%Y%m%d%H%M%S')}.png",
            "gcm_img": f"gcm_encrypted_{datetime.now().strftime('%Y%m%d%H%M%S')}.png",
            "comparison_img": comparison_filename_base
        }
    }

    # Store report and text summary in session
    session['report_data'] = report_content
    session['report_text'] = generate_text_report(report_content)

    return render_template('report.html',
                           report=session['report_text'],
                           results=results,
                           **report_content['visuals'])

@app.route('/download/<filename>')
@login_required
def download(filename):
    """
    Secure route to download encrypted image files.
    - Sanitizes filename to avoid path traversal.
    - Verifies file existence.
    - Serves the file for download as a PNG.
    """
    filename = sanitize_filename_path(filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(filepath):
        flash('⚠️ File not found', 'error')
        return redirect(url_for('secret_room'))

    try:
        return send_file(
            filepath,
            as_attachment=True,
            download_name=f"encrypted_{filename}",
            mimetype='image/png'
        )
    except Exception as e:
        logger.error(f"Error downloading file {filename}: {str(e)}")
        flash('⚠️ Error downloading file', 'error')
        return redirect(url_for('secret_room'))


@app.route('/download_report/<format>')
def download_report(format):
    """
    Downloads the generated encryption analysis report in various formats.
    Supports: PDF, Word (.docx), and plain text (.txt)
    Requires session-stored report data from /compare.
    """
    if 'report_data' not in session:
        flash('⚠️ No report available to download', 'error')
        return redirect(url_for('upload'))

    report_data = session['report_data']
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename_base = f"AES_Report_{timestamp}"

    try:
        if format == 'pdf':
            filename = f"{filename_base}.pdf"
            buffer = BytesIO()
            if generate_pdf_report(report_data, buffer):
                buffer.seek(0)
                return send_file(
                    buffer,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
            else:
                flash('⚠️ Failed to generate PDF report. Please try again.', 'error')
                return redirect(url_for('compare'))

        elif format == 'word':
            filename = f"{filename_base}.docx"
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if generate_word_report(report_data, temp_path):
                return send_file(
                    temp_path,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

        elif format == 'text':
            filename = f"{filename_base}.txt"
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(generate_text_report(report_data))

            return send_file(
                temp_path,
                as_attachment=True,
                download_name=filename,
                mimetype='text/plain'
            )

        else:
            flash('⚠️ Invalid report format requested', 'error')
            return redirect(url_for('compare'))

    except Exception as e:
        logger.error(f"Error generating {format} report: {str(e)}", exc_info=True)
        flash(f'⚠️ Failed to generate {format} report. Please try again.', 'error')
        return redirect(url_for('compare'))


@app.route('/secret_room')
@login_required
def secret_room():
    """
    Displays the user's uploaded images and their encryption status.
    - Pulls records from the database.
    - Uses UploadForm in case user wants to re-upload quickly.
    """
    form = UploadForm()
    user_id = get_user_id(session['username'])
    if not user_id:
        flash('⚠️ User not found', 'error')
        return redirect(url_for('upload'))

    try:
        with db_manager.get_db() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT filename, mode, created_at 
                FROM uploads 
                WHERE user_id = ? 
                ORDER BY created_at DESC
            ''', (user_id,))
            images = cursor.fetchall()

        return render_template('secret_room.html', images=images, form=form)

    except sqlite3.Error as e:
        logger.error(f"Database error during secret room access: {e}")
        flash('⚠️ Database error occurred. Please try again.', 'error')
        return redirect(url_for('upload'))


@app.route('/delete/<filename>')
@login_required
def delete_file(filename):
    """
    Deletes an uploaded image and removes it from the database.
    - Sanitizes input.
    - Removes both from filesystem and DB.
    """
    filename = sanitize_filename_path(filename)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    try:
        if os.path.exists(path):
            os.remove(path)

            # Remove database record
            with db_manager.get_db() as conn:
                conn.execute('DELETE FROM uploads WHERE filename = ?', (filename,))
            flash('🗑️ File deleted successfully', 'success')
        else:
            flash('⚠️ File not found', 'error')

    except Exception as e:
        logger.error(f"Error deleting file {filename}: {str(e)}")
        flash(f'⚠️ Error deleting file: {str(e)}', 'error')

    return redirect(url_for('secret_room'))


@app.route('/')
def index():
    """
    Homepage route.
    - Redirects to upload page if user is logged in.
    - Otherwise, displays landing/index page.
    """
    if 'username' in session:
        return redirect(url_for('upload'))
    return render_template('index.html')



# Initialize DB on startup
with app.app_context():
    init_db() 


# Page for sending email 
def send_email(self):
    """Render email sending page (not integrated)."""
    return render_template('index.html')


# Run Flask App
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
